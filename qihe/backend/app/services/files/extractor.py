from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt"}


class TextExtractionError(Exception):
    pass


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _normalize_text(_extract_pdf(path))
        if suffix == ".docx":
            return _normalize_text(_extract_docx(path))
        if suffix == ".txt":
            return _normalize_text(_extract_txt(path))
    except Exception as exc:
        raise TextExtractionError("文件文本抽取失败") from exc

    raise TextExtractionError("不支持的文件类型")


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraphs + table_cells)


def _extract_txt(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line).strip()
