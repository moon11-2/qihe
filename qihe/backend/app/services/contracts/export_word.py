from io import BytesIO
from typing import Any

from docx import Document

from app.models.contracts import ContractExportRequest


def export_contract_word(request: ContractExportRequest) -> bytes:
    document = Document()
    document.add_heading(request.title, level=1)

    export_type = _normalize_type(request.type)
    if export_type == "review":
        _add_review(document, request.payload)
    else:
        _add_generate(document, request.payload)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _normalize_type(export_type: str) -> str:
    if export_type == "review_result":
        return "review"
    if export_type == "generate_result":
        return "generate"
    return export_type


def _add_review(document: Document, payload: dict[str, Any]) -> None:
    review = payload.get("review_result") if "review_result" in payload else payload

    document.add_heading("审查摘要", level=2)
    document.add_paragraph(str(review.get("summary") or ""))
    if review.get("review_basis"):
        document.add_paragraph(f"审查依据：{review.get('review_basis')}")
    document.add_paragraph(f"整体风险：{review.get('risk_level', 'unknown')}")
    if review.get("score") is not None:
        document.add_paragraph(f"评分：{review.get('score')}")

    parties = review.get("parties") or {}
    if parties:
        document.add_heading("主体信息", level=2)
        for key, value in parties.items():
            document.add_paragraph(f"{key}：{value}")

    risk_items = _review_items(review)
    document.add_heading("风险项", level=2)
    if not risk_items:
        document.add_paragraph("未发现明显风险项。")
        return

    for item in risk_items:
        document.add_heading(str(_item_value(item, "risk_title", "title") or "未命名风险"), level=3)
        document.add_paragraph(f"风险等级：{_item_value(item, 'risk_level', 'level') or 'unknown'}")
        description = _item_value(item, "risk_analysis", "description")
        if description:
            document.add_paragraph(str(description))
        if item.get("clause"):
            document.add_paragraph(f"涉及条款：{item['clause']}")
        original_excerpt = _item_value(item, "original_excerpt", "excerpt")
        if original_excerpt:
            document.add_paragraph(f"原文摘录：{_truncate(str(original_excerpt), 500)}")
        start_offset = item.get("start_offset")
        end_offset = item.get("end_offset")
        if start_offset is not None and end_offset is not None:
            document.add_paragraph(f"原文位置：{start_offset}-{end_offset}")
        suggestion = _item_value(item, "revision_suggestion", "suggestion")
        if suggestion:
            document.add_paragraph(f"修订建议：{suggestion}")
        replacement = item.get("suggested_replacement")
        if replacement:
            document.add_paragraph(f"建议替换条款：{replacement}")
        legal_basis = item.get("legal_basis") or []
        if legal_basis:
            document.add_paragraph("法条依据：" + "；".join(str(value) for value in legal_basis))


def _add_generate(document: Document, payload: dict[str, Any]) -> None:
    generated = payload.get("generate_result") if "generate_result" in payload else payload

    document.add_heading("合同草案", level=2)
    draft = str(generated.get("draft") or "")
    for paragraph in draft.split("\n"):
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())

    missing_fields = generated.get("missing_fields") or []
    if missing_fields:
        document.add_heading("待补充字段", level=2)
        for field in missing_fields:
            document.add_paragraph(str(field), style="List Bullet")

    checklist = generated.get("pre_sign_checklist") or []
    if checklist:
        document.add_heading("签署前清单", level=2)
        for item in checklist:
            document.add_paragraph(str(item), style="List Bullet")

    notes = generated.get("notes") or []
    if notes:
        document.add_heading("说明", level=2)
        for item in notes:
            document.add_paragraph(str(item), style="List Bullet")


def _item_value(item: dict[str, Any], *keys: str) -> Any:
    for key in keys:
        if item.get(key):
            return item[key]
    return None


def _review_items(review: dict[str, Any]) -> list[dict[str, Any]]:
    clause_reviews = review.get("clause_reviews")
    if isinstance(clause_reviews, list) and clause_reviews:
        return clause_reviews
    risk_items = review.get("risk_items")
    if isinstance(risk_items, list):
        return risk_items
    return []


def _truncate(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return f"{text[:limit]}..."
