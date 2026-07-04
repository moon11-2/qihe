from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.core.config import settings
from app.main import create_app
from app.services.files import storage


class FailingProvider:
    async def chat_json(self, messages: list[dict[str, str]], schema_name: str) -> dict:
        raise RuntimeError("force keyword fallback")

    async def chat(self, messages: list[dict[str, str]]) -> str:
        raise RuntimeError("force chat fallback")


def _client(tmp_path: Path, monkeypatch) -> TestClient:
    monkeypatch.setattr(storage, "UPLOAD_DIR", tmp_path)
    return TestClient(create_app())


def _pdf_with_text(text: str) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=300)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    stream = DecodedStreamObject()
    stream.set_data(f"BT /F1 12 Tf 50 250 Td ({text}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(stream)

    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def test_upload_supported_file_types(tmp_path: Path, monkeypatch) -> None:
    client = _client(tmp_path, monkeypatch)

    txt_response = client.post(
        "/api/files/upload",
        files={"file": ("contract.txt", b"hello contract", "text/plain")},
    )
    assert txt_response.status_code == 200
    assert txt_response.json()["status"] == "accepted"
    assert txt_response.json()["char_count"] > 0

    docx_buffer = BytesIO()
    document = Document()
    document.add_paragraph("DOCX contract text")
    document.save(docx_buffer)
    docx_response = client.post(
        "/api/files/upload",
        files={
            "file": (
                "contract.docx",
                docx_buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert docx_response.status_code == 200
    assert "DOCX contract text" in docx_response.json()["text_preview"]

    pdf_response = client.post(
        "/api/files/upload",
        files={"file": ("contract.pdf", _pdf_with_text("PDF contract text"), "application/pdf")},
    )
    assert pdf_response.status_code == 200
    assert pdf_response.json()["filename"] == "contract.pdf"
    assert "PDF contract text" in pdf_response.json()["text_preview"]


def test_upload_rejects_images_and_large_files(tmp_path: Path, monkeypatch) -> None:
    client = _client(tmp_path, monkeypatch)

    image_response = client.post(
        "/api/files/upload",
        files={"file": ("photo.png", b"not accepted", "image/png")},
    )
    assert image_response.status_code == 400
    assert image_response.json()["error"]["code"] == "unsupported_file_type"

    monkeypatch.setattr(settings, "max_upload_mb", 0)
    large_response = client.post(
        "/api/files/upload",
        files={"file": ("contract.txt", b"x", "text/plain")},
    )
    assert large_response.status_code == 413
    assert large_response.json()["error"]["code"] == "file_too_large"


def test_upload_rejects_empty_extracted_text(tmp_path: Path, monkeypatch) -> None:
    client = _client(tmp_path, monkeypatch)

    response = client.post(
        "/api/files/upload",
        files={"file": ("empty.txt", b"   \n\t", "text/plain")},
    )

    assert response.status_code == 400
    assert response.json()["error"]["code"] == "empty_file_text"


def test_chat_response_shapes(monkeypatch) -> None:
    monkeypatch.setattr("app.services.chat.create_qwen_provider", lambda: FailingProvider())
    client = TestClient(create_app())

    review_cases = [
        "帮我审查这份合同风险",
        "帮我看下这个合同",
        "帮我把这个协议过一遍",
        "帮我把关这份条款",
    ]
    for text in review_cases:
        route_response = client.post(
            "/api/chat",
            json={"messages": [{"role": "user", "content": text}]},
        )
        assert route_response.status_code == 200
        assert route_response.json()["type"] == "route"
        assert route_response.json()["route"] == "review"

    generate_response = client.post(
        "/api/chat",
        json={"messages": [{"role": "user", "content": "帮我写一份租房合同"}]},
    )
    assert generate_response.status_code == 200
    assert generate_response.json()["type"] == "route"
    assert generate_response.json()["route"] == "generate"

    mixed_response = client.post(
        "/api/chat",
        json={"messages": [{"role": "user", "content": "我不确定，是看合同还是写合同"}]},
    )
    assert mixed_response.status_code == 200
    assert mixed_response.json()["type"] == "need_input"
    assert mixed_response.json()["options"] == ["review", "generate"]

    consult_response = client.post(
        "/api/chat",
        json={"messages": [{"role": "user", "content": "合同审查流程是什么"}]},
    )
    assert consult_response.status_code == 200
    assert consult_response.json()["type"] == "chat"

    need_input_response = client.post(
        "/api/chat",
        json={"messages": [{"role": "user", "content": "合同"}]},
    )
    assert need_input_response.status_code == 200
    assert need_input_response.json()["type"] == "need_input"
    assert need_input_response.json()["options"] == ["review", "generate"]

    chat_response = client.post(
        "/api/chat",
        json={"messages": [{"role": "user", "content": "你好"}]},
    )
    assert chat_response.status_code == 200
    assert chat_response.json()["type"] == "chat"


def test_contract_run_review_and_generate_shapes() -> None:
    client = TestClient(create_app())

    review_response = client.post(
        "/api/contracts/run",
        json={
            "mode": "review",
            "text": "甲方：甲公司\n乙方：乙公司\n双方签订服务合同，金额为人民币10000元。",
        },
    )
    assert review_response.status_code == 200
    review_json = review_response.json()
    assert review_json["type"] == "review_result"
    assert review_json["review_result"]["title"] == "合同审查报告"
    assert isinstance(review_json["review_result"]["risk_items"], list)
    assert review_json["review_result"]["parties"]["party_a"] == "甲公司"
    assert review_json["review_result"]["parties"]["party_b"] == "乙公司"
    first_review = review_json["review_result"]["clause_reviews"][0]
    for key in ("clause_id", "clause_title", "original_excerpt", "start_offset", "end_offset"):
        assert key in first_review

    generate_response = client.post(
        "/api/contracts/run",
        json={
            "mode": "generate",
            "text": "生成一份服务合同",
            "metadata": {
                "contract_type": "服务合同",
                "party_a": "甲公司",
                "party_b": "乙公司",
                "amount": "10000元",
                "term": "2026年7月1日至2026年12月31日",
                "jurisdiction": "甲方所在地人民法院",
            },
        },
    )
    assert generate_response.status_code == 200
    generate_json = generate_response.json()
    assert generate_json["type"] == "generate_result"
    assert "服务合同" in generate_json["generate_result"]["draft"]

    text_only_generate_response = client.post(
        "/api/contracts/run",
        json={
            "mode": "generate",
            "text": "生成一份服务合同，甲方甲公司，乙方乙公司，金额10000元，期限6个月。",
        },
    )
    assert text_only_generate_response.status_code == 200
    text_only_generate_json = text_only_generate_response.json()
    assert "甲公司" in text_only_generate_json["generate_result"]["draft"]
    assert "10000元" in text_only_generate_json["generate_result"]["draft"]


def test_contract_generate_uses_metadata_without_treating_empty_values_as_facts() -> None:
    client = TestClient(create_app())
    response = client.post(
        "/api/contracts/run",
        json={
            "mode": "generate",
            "text": "帮我起草一份协议",
            "metadata": {
                "contract_type": "租赁合同",
                "user_role": "普通用户",
                "my_identity": "承租方",
                "special_terms": "押金在退租验收后7日内无息退还",
                "amount": "",
                "jurisdiction": "无",
            },
        },
    )

    assert response.status_code == 200
    result = response.json()["generate_result"]
    assert "租赁合同" in result["title"] or "租赁合同" in result["draft"]
    assert "承租方" in result["draft"] or any("承租方" in note for note in result["notes"])
    assert "押金在退租验收后7日内无息退还" in result["draft"]
    assert "特殊约定" not in result["missing_fields"]
    assert "合同金额" in result["missing_fields"]
    assert any("特殊约定" in item for item in result["pre_sign_checklist"])


def test_word_export_can_be_opened() -> None:
    client = TestClient(create_app())
    response = client.post(
        "/api/contracts/export/word",
        json={
            "type": "generate",
            "title": "服务合同草案",
            "payload": {
                "draft": "服务合同\n\n甲方：甲公司\n乙方：乙公司",
                "missing_fields": ["履行期限"],
                "pre_sign_checklist": ["核对主体信息"],
            },
        },
    )

    assert response.status_code == 200
    assert response.content.startswith(b"PK")
    exported = Document(BytesIO(response.content))
    paragraphs = [paragraph.text for paragraph in exported.paragraphs]
    assert "服务合同草案" in paragraphs
    assert "履行期限" in paragraphs
